#!/usr/bin/env python3

""" Cheap Pie Module to export register descriptio into .docx file """

# this file is part of cheap_pie, a python tool for chip validation
# author: Marco Merlin
# email: marcomerli@gmail.com

import os
from ast import literal_eval
from docx import Document
# from docx.enum.table import WD_TABLE_ALIGNMENT
INCH2EMU = 914400


def int2hexstr(num, width):  # pylint: disable=W0613
    """ Convert integer into hex string """
    fmt = "%%0%dx" % width          # pylint: disable=C0209
    # print fmt
    string = "\"%s\" %% num " % fmt  # pylint: disable=C0209
    # print string
    return eval(string)  # pylint: disable=W0123


def hexstr2int(hexstr):
    """ Convert hex string into integer """
    return literal_eval('0x' + hexstr)


def doc_create_header(template=None):
    """ Create document header """
    if template is None:
        document = Document()
    else:
        document = Document(template)

    # document.add_heading('Document Title', 0)
    document.add_heading('Register Description', level=1)
    document.add_paragraph("")

    return document


def doc_add_regtable(doc, reg, tablestyle=None, nbits_addr=32):  # pylint: disable=R0912
    """ Add table description of a register to a document """
    # header
    doc.add_heading(reg.regname, level=3)

    # address
    doc.add_paragraph('Offset Address: 0x%s' % int2hexstr( # pylint: disable=C0209
        reg.addr, nbits_addr/4))

    # description
    if isinstance(reg.comments, str):
        doc.add_paragraph(reg.comments)
    else:
        print("Warning: Non-string comments for reg: " + reg.regname)

    # bitfields table
    table = doc.add_table(rows=1, cols=5)

    if tablestyle is not None:
        table.style = tablestyle
    # table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # configure column width
    # https://stackoverflow.com/questions/43749177/python-docx-table-column-width
    # inch2emu = 914400
    cwidths = [1.5,  0.5,    0.5,      1,            3]

    # create table header
    headers = ['Field name', 'rwu', 'Bit #', 'Reset', 'Description']
    hdr_cells = table.rows[0].cells

    for idx in range(len(headers)):  # pylint: disable=C0200
        hdr_cells[idx].text = headers[idx]
        hdr_cells[idx].width = cwidths[idx]*INCH2EMU
        hdr_cells[idx].bold = True
        # hdr_cells[idx].alignment=WD_ALIGN_PARAGRAPH.CENTER

    # add bitfields
    for field in reversed(reg.get_ordered_bitfields()):
        row_cells = table.add_row().cells
        row_cells[0].text = field.fieldname
        row_cells[0].width = cwidths[0]*INCH2EMU

        if len(field.read_write)/field.width > 1:
            row_cells[1].text = str.lower(field.read_write[0:2])
        else:
            row_cells[1].text = str.lower(field.read_write)
        row_cells[1].width = cwidths[1]*INCH2EMU

        if field.width == 1:
            row_cells[2].text = "%d" % (field.lsb)  # pylint: disable=C0209
        else:
            row_cells[2].text = "%d:%d" % ( # pylint: disable=C0209
                field.lsb + field.width - 1, field.lsb)
        row_cells[2].width = cwidths[2]*INCH2EMU

        if isinstance(field.reset, str):
            row_cells[3].text = field.reset
        else:
            resetstr = "%d\'%s" % (field.width, bin(field.reset)[ # pylint: disable=C0209
                                   1:])
            row_cells[3].text = resetstr
        row_cells[3].width = cwidths[3]*INCH2EMU

        if isinstance(reg.comments, str):
            row_cells[4].text = field.comments
            row_cells[4].width = cwidths[4]*INCH2EMU
            # print(field.comments)
        else:
            print(f"Warning: Non-string comments for field: \
                  {field.fieldname} in reg: {reg.regname}")

    doc.add_paragraph("")


def hal2doc(hal, fname='hal.docx', template=None, tablestyle=None, nbits_addr=32):
    """  export register description to .docx file """
    # initialize doc
    doc = doc_create_header(template)

    # loop over registers
    for reg in hal:
        # create register table
        doc_add_regtable(doc, reg, tablestyle, nbits_addr)

    # save document
    doc.save(fname)

    # check file was actually saved
    assert os.path.isfile(fname), f"ERROR: file {fname} does not exist!"


def test_hal2doc():
    """ Test function for export tool from register description to .docx """
    print('Testing hal2doc...')

    from parsers.ipxact_parse import ipxact_parse      # pylint: disable=C0413,E0401,C0415
    from cheap_pie_core.cp_cli import cp_devices_fname  # pylint: disable=C0413,E0401,C0415

    fname = cp_devices_fname("my_subblock.xml")
    hal = ipxact_parse(fname=fname)

    # convert to .docx
    hal2doc(hal)


if __name__ == '__main__':
    test_hal2doc()
