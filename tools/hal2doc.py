#!/usr/bin/python3

## this file is part of cheap_pie, a python tool for chip validation
## author: Marco Merlin
## email: marcomerli@gmail.com

from docx import Document
# from docx.enum.table import WD_TABLE_ALIGNMENT
from operator import attrgetter
# from cbitfield import cp_bitfield

import sys
import os.path
sys.path.append( os.path.join(os.path.dirname(__file__), '..') )

from cheap_pie_core.cbitfield import cp_bitfield

def reg_add_reserved_bitfields(fields,regwidth=32):
    next_lsb = 0
    outfields = []

    if isinstance(fields,list):
        if len(fields) > 0:
            for idx in reversed(range(len(fields))):
            # for idx in range(len(fields)):
                field = fields[idx]
                outfields.insert(0,field)
                #outfields.append(field)

                if next_lsb < field.lsb:
                    # print("Add Reserved field!")
                    newfield=cp_bitfield("Reserved",0,field.regname,field.lsb-next_lsb,next_lsb,"Reserved")
                    outfields.insert(1,newfield)
                    # outfields.append(newfield)

                # print("LSB: %d, WIDTH: %d" % ( field.lsb, field.width ) )
                next_lsb = field.lsb + field.width

            # check if register is filled up to full width    
            if next_lsb < regwidth:
                newfield=cp_bitfield("Reserved",0,fields[0].regname,regwidth-next_lsb,next_lsb,"Reserved")
                outfields.insert(0,newfield)
                # outfields.append(newfield)

    return outfields  

def int2hexstr(n,width):
    fmt = "%%0%dx" % width
    # print fmt 
    # string = "%0256x" % n
    # return string
    # string = "\"%0256x\" % n"
    string = "\"%s\" %% n " % fmt
    # print string    
    return eval(string)

def inch2emu():
    return 914400

def doc_create_header(template=None):
    if template is None:
        document = Document()
    else :
        document = Document(template)

    # document.add_heading('Document Title', 0)
    document.add_heading('Register Description', level=1)
    document.add_paragraph("")

    return document

def doc_add_regtable(doc,reg,tablestyle=None,nbits_addr=32):
    doc.add_heading(reg.regname, level=3)
    
    # address
    # doc.add_paragraph('Offset Address: 0x%d' % reg.addr)
    doc.add_paragraph('Offset Address: 0x%s' % int2hexstr(reg.addr,nbits_addr/4))
    # description

    if isinstance(reg.comments,str):             
        doc.add_paragraph(reg.comments)
    else:        
        Warning("Non-string comments for reg: " + reg.regname)

    # bitfields table
    table = doc.add_table(rows=1, cols=5)

    if tablestyle is not None:
        table.style = tablestyle

    # table.alignment = WD_TABLE_ALIGNMENT.CENTER
   
    # configure column width
    # https://stackoverflow.com/questions/43749177/python-docx-table-column-width
    # inch2emu = 914400
    cwidths = [         1.5,  0.5,    0.5,      1,            3]

    # create table header
    headers = ['Field name','rwu','Bit #','Reset','Description']
    hdr_cells = table.rows[0].cells

    for idx in range(len(headers)):
        hdr_cells[idx].text  = headers[idx]
        hdr_cells[idx].width = cwidths[idx]*inch2emu()
        hdr_cells[idx].bold  = True
        # hdr_cells[idx].alignment=WD_ALIGN_PARAGRAPH.CENTER

    # add bitfields
    for field in reversed(reg.bitfields):
        row_cells = table.add_row().cells
        row_cells[0].text = field.fieldname
        row_cells[0].width = cwidths[0]*inch2emu()

        if len(field.rw)/field.width > 1:
            row_cells[1].text = str.lower(field.rw[0:2])
        else:
            row_cells[1].text = str.lower(field.rw)
        row_cells[1].width = cwidths[1]*inch2emu()

        if field.width == 1:
            row_cells[2].text = "%d" % ( field.lsb )
        else:
            row_cells[2].text = "%d:%d" % ( field.lsb + field.width -1, field.lsb)
        row_cells[2].width = cwidths[2]*inch2emu()

        if isinstance(field.reset,str):
            row_cells[3].text = field.reset
        else:
            resetstr = "%d\'%s" % (field.width, bin(field.reset)[1:])
            row_cells[3].text = resetstr
        row_cells[3].width = cwidths[3]*inch2emu()

        if isinstance(reg.comments,str):             
            row_cells[4].text = field.comments
            row_cells[4].width = cwidths[4]*inch2emu()
            # print(field.comments)
        else:        
            Warning("Non-string comments for field: %s in reg: %s" % (field.fieldname, reg.regname))

    doc.add_paragraph("")

def hal2doc(hal,fname='hal.docx',template=None,tablestyle=None,nbits_addr=32):
    
    # initialize doc
    doc = doc_create_header(template)

    # loop over registers
    for reg in hal:                
        # print(reg.regname)

        # convert bitfields namedtuple into list
        bitfields = list(reg.bitfields)

        # sort by lsb
        bitfields=sorted(bitfields,key=attrgetter('lsb'))
        bitfields.reverse()

        # add reserved bitfields
        reg.bitfields = reg_add_reserved_bitfields(bitfields)

        # create register table
        doc_add_regtable(doc,reg,tablestyle,nbits_addr)

    # save document
    doc.save(fname)

def test_hal2doc():
    print('Testing hal2doc...')
    if False:
        from parsers.svd_parse import svd_parse
        # from parsers.svd_parse_repo import svd_parse
        hal = svd_parse(fname="./devices/QN908XC.svd")
    else:
        from parsers.ipxact_parse import ipxact_parse
        # from parsers.svd_parse_repo import svd_parse
        hal = ipxact_parse(fname="./devices/my_subblock.xml")
    # convert to .docx
    hal2doc(hal)

if __name__ == '__main__':
    test_hal2doc()
    pass